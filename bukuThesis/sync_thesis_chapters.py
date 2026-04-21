#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/home/hduser/jupyter/gust/RisetEU/bukuThesis")
BAHAN = ROOT / "bahan"

METODOLOGI = BAHAN / "progress_bab-metodologiPenelitian_updated.docx"
BAB_I = BAHAN / "progress_bab-implementasi_dan_hasil_pembahasan_updated.docx"
PROGRESS_SKENARIO = BAHAN / "progres_BAba Progress PEnelitian dan Skenario pengujian_updated.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_first(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Tidak menemukan paragraf dengan teks: {needle}")


def contains(doc: Document, needle: str) -> bool:
    return any(needle in p.text for p in doc.paragraphs)


def sync_bab_i() -> None:
    doc = Document(BAB_I)

    if not contains(doc, "4.1 Analisis Konvergensi, Overfitting, dan Diagnostik Residual Model"):
        anchor = find_first(doc, "5. Pembahasan Hasil")
        blocks = [
            (
                "4.1 Analisis Konvergensi, Overfitting, dan Diagnostik Residual Model",
                "Heading 3",
            ),
            (
                "Analisis tambahan dilakukan untuk memastikan bahwa perbedaan performa antar model tidak hanya dibaca dari MAE(pips), tetapi juga dari perilaku pelatihan dan kualitas residual forecast. Pada FLF-LSTM dan FLF-BiLSTM, analisis dilakukan melalui kurva loss dan val_loss, sedangkan pada ARIMA analisis yang setara dilakukan melalui diagnostik residual, AIC/BIC, dan uji autokorelasi residual. Adapun grafik gradient loss pada artefak repo diperlakukan sebagai proxy perubahan loss per epoch (dLoss), sehingga fungsinya adalah membaca kecepatan konvergensi dan plateau, bukan gradient parameter model.",
                "Normal",
            ),
            (
                "Pada FLF-LSTM, rata-rata epoch berjalan selama lima fold terakhir adalah 53.8, dengan rata-rata loss_end = 0.000485, val_end = 0.000547, dan gap akhir val_loss - loss = 0.000062. Validation loss umumnya mencapai titik minimum pada bagian akhir training, yaitu sekitar 95.89% dari total epoch, dan hanya menunjukkan pola overfitting ringan pada 3 dari 5 fold. Temuan ini menunjukkan bahwa FLF-LSTM berhasil konvergen, tidak underfit, dan memiliki generalisasi yang relatif stabil pada skenario walk-forward 72 bulan / 1 bulan.",
                "Normal",
            ),
            (
                "Pada FLF-BiLSTM, rata-rata epoch berjalan adalah 37.6, dengan loss_end = 0.000498, val_end = 0.000597, dan gap akhir 0.000099. Validation loss cenderung mencapai titik minimum lebih awal, yaitu sekitar 89.30% dari total epoch, lalu pada beberapa fold meningkat kembali ketika training loss masih menurun. Pola tersebut muncul pada 3 dari 5 fold, sehingga FLF-BiLSTM dapat dinyatakan tetap konvergen dan tidak underfit, tetapi lebih rentan terhadap overfitting dibanding FLF-LSTM.",
                "Normal",
            ),
            (
                "Untuk baseline ARIMA, analisis tidak dilakukan melalui kurva loss berbasis epoch karena mekanisme modelnya berbeda dari neural network. Evaluasi yang lebih tepat adalah diagnostik residual forecast. Pada lima fold terakhir, ARIMA menghasilkan combined close residual bias = 0.7157 pips, combined residual std = 22.3251 pips, lag-1 autocorrelation = -0.0018, Ljung-Box p-value lag 10 = 0.0566, dan lag 20 = 0.0500. Hasil ini menunjukkan bahwa residual forecast ARIMA relatif stabil dan tidak memperlihatkan bukti kuat autokorelasi serial pada lag pendek, walaupun pada horizon gabungan yang lebih panjang masih terdapat indikasi lemah bahwa sebagian struktur temporal belum sepenuhnya hilang.",
                "Normal",
            ),
            (
                "Secara komparatif, analisis ini mendukung hasil MAE(pips) sebelumnya. FLF-LSTM bukan hanya memberikan error agregat terendah, tetapi juga menunjukkan dinamika pelatihan yang paling sehat dan stabil. FLF-BiLSTM masih kompetitif dan tetap mengungguli ARIMA, namun generalisasinya lebih cepat jenuh. Sementara itu, ARIMA tetap valid sebagai baseline statistik yang stabil dan mudah diinterpretasikan, tetapi belum mampu menyamai kualitas generalisasi model FLF pada skenario evaluasi utama.",
                "Normal",
            ),
        ]
        prev = anchor
        for text, style in reversed(blocks):
            prev = insert_paragraph_before(prev, text, style)

    replacements = {
        "hasil perbandingan FLF-LSTM dan FLF-BiLSTM pada results/lstm_vs_bilstm_wf72_test1_last5_comparison.html,":
            "1. hasil perbandingan tiga model pada results/comparison/comparison_models_wf72_test1_last5.html,",
        "1. hasil perbandingan tiga model pada results/model_comparison_wf72_test1_last5.html,":
            "1. hasil perbandingan tiga model pada results/comparison/comparison_models_wf72_test1_last5.html,",
        "1. hasil perbandingan tiga model pada results/comparison/model_comparison_wf72_test1_last5.html,":
            "1. hasil perbandingan tiga model pada results/comparison/comparison_models_wf72_test1_last5.html,",
        "hasil tuning FLF-LSTM dan FLF-BiLSTM pada bukuThesis/optimasi_hyperparameter_flf_bilstm_lstm_eurusd.md, dan":
            "2. hasil tuning FLF-LSTM dan FLF-BiLSTM pada bukuThesis/optimasi_hyperparameter_flf_bilstm_lstm_eurusd.md,",
        "baseline ARIMA pada Arima/result/arima_wf72_test1_last5/.":
            "3. baseline ARIMA pada Arima/result/arima_wf72_test1_last5/, dan",
    }
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt in replacements:
            p.text = replacements[txt]

    if not contains(doc, "4. analisis konvergensi dan diagnostik model pada bukuThesis/bahan/analisis_konvergensi_dan_diagnostik_model_wf72_test1_last5.md."):
        anchor = find_first(doc, "3. baseline ARIMA pada Arima/result/arima_wf72_test1_last5/, dan")
        insert_paragraph_after(
            anchor,
            "4. analisis konvergensi dan diagnostik model pada bukuThesis/bahan/analisis_konvergensi_dan_diagnostik_model_wf72_test1_last5.md.",
            "Normal",
        )

    doc.save(BAB_I)


def sync_metodologi() -> None:
    doc = Document(METODOLOGI)
    if not contains(doc, "Dalam penelitian ini, horizon testing ditetapkan tetap selama 1 bulan"):
        anchor = find_first(doc, "Untuk menjaga fairness antar model, skenario evaluasi komparatif utama menggunakan walk-forward fixed dengan 72 bulan data latih dan 1 bulan data uji.")
        p1 = insert_paragraph_after(
            anchor,
            "Dalam penelitian ini, horizon testing ditetapkan tetap selama 1 bulan. Pemilihan horizon ini didasarkan pada kebutuhan untuk menjaga konsistensi protokol evaluasi antar model, mempertahankan relevansi operasional pada pasar forex, serta menyediakan jumlah sampel uji yang cukup representatif pada timeframe 4H. Dengan horizon 1 bulan, evaluasi masih berfokus pada kondisi pasar terbaru tanpa membuat ukuran sampel uji terlalu pendek seperti 14 hari atau terlalu panjang seperti 3 bulan. Oleh karena itu, horizon 1 bulan diperlakukan sebagai skenario evaluasi utama, sedangkan horizon lain yang pernah dicoba hanya diposisikan sebagai konteks eksploratif dan bukan basis utama penarikan kesimpulan penelitian.",
            "Normal",
        )
        insert_paragraph_after(
            p1,
            "Sementara itu, panjang jendela training ditetapkan sebesar 72 bulan dan diperlakukan sebagai keputusan desain evaluasi yang rasional, bukan sebagai hasil optimasi yang mengklaim nilai terbaik secara absolut. Pemilihan 72 bulan dimaksudkan untuk menyediakan konteks historis yang cukup panjang bagi model, menjaga kecukupan sampel latih, dan memastikan protokol evaluasi yang seragam pada seluruh model yang dibandingkan. Dengan demikian, penggunaan walk-forward fixed 72 bulan training dan 1 bulan testing diposisikan sebagai kerangka evaluasi utama yang defensif secara metodologis, meskipun penelitian ini tidak mengklaim bahwa 72 bulan telah dibuktikan lebih unggul daripada seluruh alternatif panjang jendela train lainnya.",
            "Normal",
        )

    if not contains(doc, "Sebagai analisis pendukung, model FLF dievaluasi pula melalui kurva loss dan val_loss"):
        anchor = find_first(doc, "Pada P5 – Output & Evaluation")
        insert_paragraph_after(
            anchor,
            "Sebagai analisis pendukung, model FLF dievaluasi pula melalui kurva loss dan val_loss untuk menilai konvergensi, indikasi overfitting, dan stabilitas generalisasi. Grafik gradient loss pada repo diperlakukan sebagai proxy perubahan loss per epoch (dLoss), bukan gradient parameter model. Untuk baseline ARIMA, analisis pendukung dilakukan melalui diagnostik residual forecast, AIC/BIC, dan uji autokorelasi residual seperti Ljung-Box, sehingga setiap model memiliki bentuk evaluasi diagnostik yang sesuai dengan mekanisme pemodelannya.",
            "Normal",
        )
    doc.save(METODOLOGI)


def sync_progress_skenario() -> None:
    doc = Document(PROGRESS_SKENARIO)

    if not contains(doc, "Artefak komparatif utama saat ini sudah mencakup report tiga model"):
        anchor = find_first(doc, "Pada jalur FLF, optimasi hyperparameter baseline OHLC-only telah selesai dijalankan")
        insert_paragraph_after(
            anchor,
            "Artefak komparatif utama saat ini sudah mencakup report tiga model pada skenario yang sama, yaitu ARIMA vs FLF-LSTM vs FLF-BiLSTM untuk fold 17-21. Pada skenario tersebut, mean MAE(pips) masing-masing model adalah 15.3941 untuk ARIMA, 11.3646 untuk FLF-LSTM, dan 12.3993 untuk FLF-BiLSTM. Hasil ini menempatkan FLF-LSTM sebagai model dengan performa agregat terbaik pada rezim pasar terbaru yang diuji.",
            "Normal",
        )

    if not contains(doc, "analisis konvergensi dan overfitting untuk FLF-LSTM dan FLF-BiLSTM"):
        anchor = find_first(doc, "Penggunaan MAE(pips) dimaksudkan untuk mempertahankan proporsi error asli")
        insert_paragraph_after(
            anchor,
            "Selain metrik agregat tersebut, skenario pengujian pada tahap progress ini juga dilengkapi analisis pendukung berupa analisis konvergensi dan overfitting untuk FLF-LSTM dan FLF-BiLSTM, serta diagnostik residual forecast untuk baseline ARIMA. Analisis tambahan ini tidak menggantikan MAE(pips) sebagai metrik utama, tetapi membantu memastikan bahwa perbedaan error antar model konsisten dengan perilaku pelatihan dan kualitas residual masing-masing model.",
            "Normal",
        )

    if not contains(doc, "Horizon uji 1 bulan dipilih sebagai fokus utama"):
        anchor = find_first(doc, "Ringkasan hasil komparatif utama menggunakan lima fold terakhir, yaitu fold 17 sampai fold 21.")
        p1 = insert_paragraph_after(
            anchor,
            "Horizon uji 1 bulan dipilih sebagai fokus utama karena memberikan keseimbangan yang lebih baik antara relevansi operasional pasar forex dan kecukupan sampel evaluasi pada timeframe 4H. Dengan horizon ini, model tetap diuji pada kondisi pasar terbaru, tetapi jumlah observasi uji masih cukup untuk menghasilkan perbandingan yang bermakna antar model.",
            "Normal",
        )
        insert_paragraph_after(
            p1,
            "Jendela training 72 bulan dipilih sebagai keputusan desain evaluasi untuk menyediakan konteks historis yang cukup panjang dan menjaga fairness protokol antar model. Pilihan ini tidak dimaknai sebagai klaim bahwa 72 bulan adalah nilai train window terbaik secara absolut, melainkan sebagai konfigurasi evaluasi utama yang rasional dan konsisten untuk ARIMA, FLF-LSTM, dan FLF-BiLSTM.",
            "Normal",
        )

    doc.save(PROGRESS_SKENARIO)


def main() -> None:
    sync_metodologi()
    sync_bab_i()
    sync_progress_skenario()
    print(METODOLOGI)
    print(BAB_I)
    print(PROGRESS_SKENARIO)


if __name__ == "__main__":
    main()
